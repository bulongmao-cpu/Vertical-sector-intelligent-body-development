from toolbox import update_ui, CatchException, report_exception, write_history_to_file, promote_file_to_downloadzone
import pandas as pd
import numpy as np
import os, glob
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import GridSearchCV
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
from sklearn.multioutput import MultiOutputRegressor
from xgboost import XGBRegressor
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive
import markdown2
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

def save_markdown_to_styled_docx(md_text, path, image_paths=None):
    html = markdown2.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    for elem in soup.descendants:
        if elem.name in ["h1", "h2", "p", "li"]:
            text = elem.text.strip()
            if not text: continue
            if elem.name == "h1":
                p = doc.add_heading(text, level=1)
            elif elem.name == "h2":
                p = doc.add_heading(text, level=2)
            elif elem.name == "li":
                p = doc.add_paragraph(text, style="List Bullet")
            else:
                p = doc.add_paragraph(text)
            run = p.runs[0]
            run.font.name = '宋体'
            p.style.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
    if image_paths:
        for img in image_paths:
            doc.add_paragraph(f"图：{os.path.basename(img)}")
            doc.add_picture(img, width=Inches(5.5))
    doc.save(path)

def run_full_xgboost_model(df: pd.DataFrame, save_dir: str):
    feature_columns = [
        "is_center", "building_area", "land_area", "green_area",
        "plot_ratio", "building_count", "dorm_count", "plot", "year"
    ]
    target_columns = [
        "security_cost", "cleaning_cost", "landscaping_cost",
        "maintenance_cost", "dorm_admin_cost", "total_cost"
    ]
    df = df.dropna(subset=feature_columns + target_columns)
    df["dorm_admin_cost"] += df["dorm_count"] * 0.5
    X = df[feature_columns]
    y = df[target_columns]

    r2_table, mae_table, rmse_table = {}, {}, {}
    feature_importance_df = pd.DataFrame(index=feature_columns)
    images = []

    for i, target in enumerate(target_columns):
        model = XGBRegressor(objective='reg:squarederror')
        grid = GridSearchCV(model, {
            "n_estimators": [30],
            "max_depth": [3],
            "learning_rate": [0.1]
        }, cv=3, scoring="r2")
        grid.fit(X, y[target])
        best_model = grid.best_estimator_
        preds = best_model.predict(X)
        r2_table[target] = round(r2_score(y[target], preds), 3)
        mae_table[target] = round(mean_absolute_error(y[target], preds), 3)
        rmse_table[target] = round(mean_squared_error(y[target], preds, squared=False), 3)

        importances = best_model.feature_importances_
        feature_importance_df[target] = importances

        # 保存图
        plt.figure(figsize=(8, 4))
        sns.barplot(x=importances, y=feature_columns)
        plt.title(f"{target} 特征重要性")
        img_path = os.path.join(save_dir, f"importance_{target}.png")
        plt.tight_layout()
        plt.savefig(img_path)
        images.append(img_path)
        plt.close()

    return {
        "r2": r2_table,
        "mae": mae_table,
        "rmse": rmse_table,
        "importance": feature_importance_df,
        "images": images
    }

@CatchException
def 成本动因分析_XGBoost_LLM完整版(main_input, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    chatbot.append(["📊 成本动因分析任务启动", "准备读取Excel数据..."])
    yield from update_ui(chatbot=chatbot, history=history)

    if os.path.isdir(main_input):
        excel_files = glob.glob(os.path.join(main_input, "*.xlsx"))
        if not excel_files:
            chatbot.append(["❌ 未找到Excel文件", main_input])
            yield from update_ui(chatbot=chatbot, history=history)
            return
        filepath = excel_files[0]
    else:
        filepath = main_input

    df = pd.read_excel(filepath)
    chatbot.append(["📈 模型运行中", "正在进行多输出回归建模..."])
    yield from update_ui(chatbot=chatbot, history=history)

    result = run_full_xgboost_model(df, os.path.dirname(filepath))

    summary = f"""
你是一个成本建模分析专家，请根据以下指标撰写正式报告段落：
R²：{result['r2']}
MAE：{result['mae']}
RMSE：{result['rmse']}
以下是部分特征重要性（DataFrame）：
{result['importance'].head().to_markdown()}
请撰写结构化报告段落，涵盖模型效果、关键变量、优化建议等。
"""

    gpt_reply = yield from request_gpt_model_in_new_thread_with_ui_alive(
        inputs=summary,
        inputs_show_user="请生成结构化动因分析报告",
        llm_kwargs=llm_kwargs,
        chatbot=chatbot,
        history=history,
        sys_prompt="你是成本分析专家，结合建模结果与图像撰写结构化中文报告。"
    )

    history.append(("最终分析报告", gpt_reply))
    res = write_history_to_file(history)
    promote_file_to_downloadzone(res, chatbot=chatbot)

    word_path = filepath.replace(".xlsx", "_成本分析报告.docx")
    save_markdown_to_styled_docx(gpt_reply, word_path, result["images"])
    promote_file_to_downloadzone(word_path, chatbot=chatbot)
    chatbot.append(["📄 Word报告生成完成", word_path])
    yield from update_ui(chatbot=chatbot, history=history)
