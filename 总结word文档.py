from toolbox import update_ui
from toolbox import CatchException, report_exception
from toolbox import write_history_to_file, promote_file_to_downloadzone
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive
fast_debug = False


def export_summary_docx(doc_title: str,
                        fragment_summaries: list,
                        final_summary: str,
                        save_dir: str = "./tmp") -> str:
    """
    导出 Word 摘要文档（中文统一宋体 SimSun，稳健不崩）。
    """
    import os, time, re, traceback
    from pathlib import Path
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    CN_FONT = "SimSun"   # 中文固定宋体
    EN_FONT = "Calibri"

    def strip_think(text: str) -> str:
        return re.sub(r"<think>.*?</think>", "", text or "", flags=re.S).strip()

    def safe_set_run_font(run, size_pt=None, bold=False):
        """仅设置字体与 eastAsia，不再触碰 rPr.lang。"""
        try:
            run.font.name = EN_FONT
            if size_pt:
                run.font.size = Pt(size_pt)
            run.bold = bool(bold)
            rpr = run._element.get_or_add_rPr()
            rfonts = getattr(rpr, "rFonts", None)
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            rfonts.set(qn('w:ascii'), EN_FONT)
            rfonts.set(qn('w:hAnsi'), EN_FONT)
            rfonts.set(qn('w:eastAsia'), CN_FONT)   # 关键：中文走宋体
        except Exception:
            # 不让异常中断
            pass

    def safe_par_spacing(p, line_spacing=1.3, before=6, after=6):
        try:
            fmt = p.paragraph_format
            fmt.line_spacing = line_spacing
            fmt.space_before = Pt(before)
            fmt.space_after = Pt(after)
            # 清除列表/编号
            for numPr in p._element.xpath(".//w:numPr"):
                numPr.getparent().remove(numPr)
        except Exception:
            pass

    # 清洗文本
    fragment_summaries = [strip_think(s) for s in (fragment_summaries or [])]
    final_summary = strip_think(final_summary or "")

    # 路径
    out_dir = Path(save_dir).absolute()
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"summary_{int(time.time())}.docx"

    # 生成文档（样式失败也不影响写入）
    try:
        doc = Document()

        # 标题
        title_para = doc.add_heading(f"文本摘要 - {doc_title}", level=1)
        safe_par_spacing(title_para, 1.25, 0, 8)
        for r in title_para.runs:
            safe_set_run_font(r, size_pt=18, bold=True)

        # 分段摘要
        if fragment_summaries:
            h2 = doc.add_heading("分段摘要", level=2)
            safe_par_spacing(h2, 1.25, 12, 6)
            for r in h2.runs:
                safe_set_run_font(r, size_pt=14, bold=True)

            for i, s in enumerate(fragment_summaries, 1):
                tag_p = doc.add_paragraph(f"【分段 {i}】")
                safe_par_spacing(tag_p, 1.3, 4, 2)
                for r in tag_p.runs:
                    safe_set_run_font(r, size_pt=11, bold=True)

                for line in str(s).splitlines():
                    p = doc.add_paragraph(line)
                    safe_par_spacing(p, 1.3, 0, 2)
                    for r in p.runs:
                        safe_set_run_font(r, size_pt=11)

        # 最终摘要
        if final_summary:
            h2 = doc.add_heading("最终摘要", level=2)
            safe_par_spacing(h2, 1.25, 12, 6)
            for r in h2.runs:
                safe_set_run_font(r, size_pt=14, bold=True)

            for line in str(final_summary).splitlines():
                p = doc.add_paragraph(line)
                safe_par_spacing(p, 1.3, 0, 2)
                for r in p.runs:
                    safe_set_run_font(r, size_pt=11)

        doc.save(str(out_path))
        return str(out_path)

    except Exception:
        # 兜底：极简版，保证一定落盘
        doc = Document()
        doc.add_heading(f"文本摘要 - {doc_title}", level=1)
        if fragment_summaries:
            doc.add_heading("分段摘要", level=2)
            for i, s in enumerate(fragment_summaries, 1):
                doc.add_paragraph(f"【分段 {i}】")
                doc.add_paragraph(str(s))
        if final_summary:
            doc.add_heading("最终摘要", level=2)
            doc.add_paragraph(str(final_summary))
        doc.save(str(out_path))
        return str(out_path)


def 解析docx(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt):
    import time, os
    # pip install python-docx 用于docx格式，跨平台
    # pip install pywin32 用于doc格式，仅支持Win平台
    for index, fp in enumerate(file_manifest):
        fragment_summaries = []
        final_summary = ""

        if fp.split(".")[-1] == "docx":
            from docx import Document
            doc = Document(fp)
            file_content = "\n".join([para.text for para in doc.paragraphs])
        else:
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.visible = False
                # 打开文件
                doc = word.Documents.Open(os.getcwd() + '/' + fp)
                # file_content = doc.Content.Text
                doc = word.ActiveDocument
                file_content = doc.Range().Text
                doc.Close()
                word.Quit()
            except:
                raise RuntimeError('请先将.doc文档转换为.docx文档。')

        # private_upload里面的文件名在解压zip后容易出现乱码（rar和7z格式正常），故可以只分析文章内容，不输入文件名
        from crazy_functions.pdf_fns.breakdown_txt import breakdown_text_to_satisfy_token_limit
        from request_llms.bridge_all import model_info
        max_token = model_info[llm_kwargs['llm_model']]['max_token']
        TOKEN_LIMIT_PER_FRAGMENT = max_token * 3 // 4
        paper_fragments = breakdown_text_to_satisfy_token_limit(txt=file_content, limit=TOKEN_LIMIT_PER_FRAGMENT, llm_model=llm_kwargs['llm_model'])
        this_paper_history = []
        for i, paper_frag in enumerate(paper_fragments):
            i_say = f'请对下面的文章片段用中文做概述，文件名是{os.path.relpath(fp, project_folder)}，文章内容是 ```{paper_frag}```'
            i_say_show_user = f'请对下面的文章片段做概述: {os.path.abspath(fp)}的第{i+1}/{len(paper_fragments)}个片段。'
            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say_show_user,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=[],
                sys_prompt="总结文章。"
            )

            # —— 新增：收集分段摘要 ——
            fragment_summaries.append(gpt_say)

            chatbot[-1] = (i_say_show_user, gpt_say)
            history.extend([i_say_show_user,gpt_say])
            this_paper_history.extend([i_say_show_user,gpt_say])

        # 已经对该文章的所有片段总结完毕，如果文章被切分了，
        if len(paper_fragments) > 1:
            i_say = f"根据以上的对话，总结文章{os.path.abspath(fp)}的主要内容。"
            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=this_paper_history,
                sys_prompt="总结文章。"
            )

            # —— 新增：最终汇总摘要 ——
            final_summary = gpt_say

            history.extend([i_say, gpt_say])
            this_paper_history.extend([i_say, gpt_say])

        else:
            # 只有一个片段时，最后一次就是最终摘要
            final_summary = fragment_summaries[-1] if fragment_summaries else ""

            # —— 新增：导出 .docx 并放到下载区 ——
        try:
            docx_path = export_summary_docx(
                doc_title=os.path.basename(fp),
                fragment_summaries=fragment_summaries,
                final_summary=final_summary,
                save_dir="./tmp"
            )
            promote_file_to_downloadzone(docx_path, chatbot=chatbot)
            chatbot.append(("📄 已生成摘要Word文档", docx_path))
        except Exception as e:
            chatbot.append(("⚠️ 生成Word文档失败", str(e)))

        res = write_history_to_file(history)
        promote_file_to_downloadzone(res, chatbot=chatbot)
        chatbot.append(("完成了吗？", res))
        yield from update_ui(chatbot=chatbot, history=history) # 刷新界面

    res = write_history_to_file(history)
    promote_file_to_downloadzone(res, chatbot=chatbot)
    chatbot.append(("所有文件都总结完成了吗？", res))
    yield from update_ui(chatbot=chatbot, history=history) # 刷新界面


@CatchException
def 总结word文档(txt, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    import glob, os

    # 基本信息：功能、贡献者
    chatbot.append([
        "函数插件功能？",
        "批量总结Word文档。注意, 如果是.doc文件, 请先转化为.docx格式。"])
    yield from update_ui(chatbot=chatbot, history=history) # 刷新界面

    # 尝试导入依赖，如果缺少依赖，则给出安装建议
    try:
        from docx import Document
    except:
        report_exception(chatbot, history,
                         a=f"解析项目: {txt}",
                         b=f"导入软件依赖失败。使用该模块需要额外依赖，安装方法```pip install --upgrade python-docx pywin32```。")
        yield from update_ui(chatbot=chatbot, history=history) # 刷新界面
        return

    # 清空历史，以免输入溢出
    history = []

    # 检测输入参数，如没有给定输入参数，直接退出
    if os.path.exists(txt):
        project_folder = txt
    else:
        if txt == "": txt = '空空如也的输入栏'
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"找不到本地项目或无权访问: {txt}")
        yield from update_ui(chatbot=chatbot, history=history) # 刷新界面
        return

    # 搜索需要处理的文件清单
    if txt.endswith('.docx') or txt.endswith('.doc'):
        file_manifest = [txt]
    else:
        file_manifest = [f for f in glob.glob(f'{project_folder}/**/*.docx', recursive=True)] + \
                        [f for f in glob.glob(f'{project_folder}/**/*.doc', recursive=True)]

    # 如果没找到任何文件
    if len(file_manifest) == 0:
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"找不到任何.docx或doc文件: {txt}")
        yield from update_ui(chatbot=chatbot, history=history) # 刷新界面
        return

    # 开始正式执行任务
    yield from 解析docx(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt)
