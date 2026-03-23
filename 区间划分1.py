from sklearn.tree import DecisionTreeRegressor
from scipy.stats import f_oneway
import mapclassify
from toolbox import update_ui, CatchException, report_exception, write_history_to_file, promote_file_to_downloadzone
import pandas as pd
import numpy as np
from sklearn.linear_model import LinearRegression
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive
import os
import glob
import matplotlib.pyplot as plt
import seaborn as sns
import markdown2
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import matplotlib
import re
from sklearn.preprocessing import StandardScaler
from collections import defaultdict, Counter
from matplotlib.font_manager import FontProperties

matplotlib.use('Agg')
matplotlib.rcParams['font.family'] = 'AR PL UKai CN'
matplotlib.rcParams['axes.unicode_minus'] = False 
matplotlib.rcParams['font.sans-serif'] = ['AR PL UKai CN']

def df_to_word_table(doc, df):
    # 添加表格，行数是数据行数+1（标题行），列数是DataFrame列数
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = 'Table Grid'  # 你可以选择其他样式
    
    # 设置标题行
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        # 设置字体（宋体，字号12）
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'AR PL UKai CN'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'AR PL UKai CN')
        run.font.size = Pt(12)

    # 填充数据行
    for row_idx in range(df.shape[0]):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx in range(df.shape[1]):
            val = df.iat[row_idx, col_idx]
            row_cells[col_idx].text = str(val)
            paragraph = row_cells[col_idx].paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'AR PL UKai CN'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'AR PL UKai CN')
            run.font.size = Pt(12)

def add_summary_table_to_doc(doc, summary):
    """
    将 summary (pandas 多级列索引DataFrame) 转为 Word 表格插入 doc。
    summary 格式示例：df.groupby(...).agg(["mean", "std", "count"])
    """
    # 先处理表头（两层列索引）
    columns = summary.columns
    # 第一层和第二层标题
    level0 = [col[0] for col in columns]
    level1 = [col[1] for col in columns]

    nrows = summary.shape[0] + 2  # 2 行表头 + 数据行
    ncols = len(columns) + 1      # 多级列 + 分段标签列

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"

    # 第一行：第一层标题，分段标签占一列
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "分段标签"
    col_pos = 1
    prev = None
    span_start = 1
    for i, val in enumerate(level0):
        hdr_cells[col_pos].text = val
        col_pos += 1

    # 第二行：第二层标题，分段标签列空白
    hdr_cells2 = table.rows[1].cells
    hdr_cells2[0].text = ""
    for i, val in enumerate(level1):
        hdr_cells2[i + 1].text = val

    # 填充数据
    for row_idx, idx in enumerate(summary.index):
        row_cells = table.rows[row_idx + 2].cells
        row_cells[0].text = str(idx)
        for col_idx, col in enumerate(columns):
            val = summary.iloc[row_idx, col_idx]
            # 格式化数字，非数字或nan直接转字符串
            if pd.isna(val):
                row_cells[col_idx + 1].text = ""
            elif isinstance(val, (int, np.integer)):
                row_cells[col_idx + 1].text = str(val)
            elif isinstance(val, (float, np.floating)):
                row_cells[col_idx + 1].text = f"{val:.4f}"
            else:
                row_cells[col_idx + 1].text = str(val)

def clean_and_prepare(df, split_col, target_cols):
    df = df.copy()
    df.columns = df.columns.str.strip()

    # 清理数值格式（去掉逗号，转成 float）
    for col in [split_col] + target_cols:
        df[col] = df[col].astype(str).str.replace(",", "").replace("nan", None)
        df[col] = pd.to_numeric(df[col], errors="coerce")

    # 去除这些列中含 NaN 的行
    df = df.dropna(subset=[split_col] + target_cols)

    cleaned_df = df.copy()
    outlier_dict = {}

    for col in target_cols:
        series = cleaned_df[col]
        Q1 = series.quantile(0.25)
        Q3 = series.quantile(0.75)
        IQR = Q3 - Q1

        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR

        outliers = series[(series < lower_bound) | (series > upper_bound)]
        outlier_dict[col] = outliers

        # 从 cleaned_df 中去除这些离群值
        cleaned_df = cleaned_df[~cleaned_df[col].isin(outliers)]

    return cleaned_df, outlier_dict

def save_markdown_to_docx(md_text, path, image_map=None, summary=None):
    html = markdown2.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    img_flag = 1
    doc = Document()
    if image_map is None:
        image_map = {}

    for elem in soup.descendants:
        if elem.name in ["h1", "h2", "p", "li"]:
            text = elem.text.strip()
            if not text:
                continue

            if elem.name == "h1":
                p = doc.add_heading(text, level=1)
            elif elem.name == "h2":
                p = doc.add_heading(text, level=2)
            elif elem.name == "li":
                p = doc.add_paragraph(text, style="List Bullet")
            else:
                p = doc.add_paragraph(text)

            run = p.runs[0]
            run.font.name = 'AR PL UKai CN'
            p.style.font.name = 'AR PL UKai CN'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'AR PL UKai CN')
            if img_flag == 1:
                for img_path in image_map.values():
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(5.5))
                img_flag=0

    if summary is not None:
        doc.add_page_break()
        doc.add_heading("分段统计汇总表", level=2)
        df_to_word_table(doc, summary)
    doc.save(path)


def segment_by_area(df, split_col, target_cols, method="tree", max_segments=None, min_ratio=0.3):

    df = df.dropna(subset=target_cols + [split_col])
    df = df.sort_values(by=split_col).reset_index(drop=True)
    N = len(df)

    best_score = -np.inf
    best_split = None
    best_labels = None

    def calc_score(groups):
        score = 0
        for col in target_cols:
            group_vals = [g[col].values for g in groups]
            if all(len(g) > 1 for g in group_vals):
                try:
                    F, p = f_oneway(*group_vals)
                    score += F  # 或 -np.log(p + 1e-8)
                except:
                    pass
        return score

    if method == "maxdiff":
        for i in range(int(N * min_ratio), int(N * (1 - min_ratio))):
            g1 = df.iloc[:i]
            g2 = df.iloc[i:]
            score = calc_score([g1, g2])
            if score > best_score:
                best_score = score
                best_split = [df.iloc[i][split_col]]
                best_labels = [0] * i + [1] * (N - i)

    elif method == "tree":
        # 用单位面积成本平均值作为监督信号
        # 使用所有目标变量的均值作为监督信号
        scaler = StandardScaler()
        target_std = pd.DataFrame(
            scaler.fit_transform(df[target_cols]),
            columns=target_cols,
            index=df.index
        )
        y = target_std.mean(axis=1).values

        X = df[[split_col]].values
        tree = DecisionTreeRegressor(max_leaf_nodes=max_segments, min_samples_leaf=int(N * min_ratio))
        tree.fit(X, y)
        labels = tree.apply(X)
        df["分段标签"] = pd.Series(labels).rank(method="dense").astype(int) - 1
        best_labels = df["分段标签"].tolist()

        # 找标签变化的位置索引
        change_idx = df.index[df["分段标签"].diff() != 0].tolist()

        # 计算边界：变化点与前一点的split_col均值
        boundaries = []
        for idx in change_idx:
            if idx > 0:
                boundary = round((df.loc[idx, split_col] + df.loc[idx - 1, split_col]) / 2, 4)
                boundaries.append(boundary)

        best_split = sorted(boundaries)

    elif method == "jenks":
        values = df[split_col].values
        for k in range(max_segments, 1, -1):
            classifier = mapclassify.NaturalBreaks(values, k=k)
            labels = classifier.yb
            counts = pd.Series(labels).value_counts(normalize=True)
            print(f"尝试 {k} 段，各段占比：\n{counts}")
            if (counts >= min_ratio).all():
                df["分段标签"] = labels
                best_labels = labels.tolist()
                best_split = classifier.bins[:k-1].tolist()
                break
        else:
            df["分段标签"] = 0
            best_labels = [0] * N
            best_split = []

    df["分段标签"] = best_labels
    boundaries = sorted(best_split)

    # 统计分析输出
    summary_raw = df.groupby("分段标签")[[split_col] + target_cols].agg(["mean", "std", "count"])

    # 分段区间标签转换（单位转成万平方米方便阅读）
    bins = [df[split_col].min()] + boundaries + [df[split_col].max()]
    interval_labels = []
    for i in range(len(bins) - 1):
        left = round(bins[i], 2)
        right = round(bins[i + 1], 2)
        interval_labels.append(f"{left}-{right}万平方米")

    # 把分段数字标签映射成区间标签
    summary_raw.index = [interval_labels[i] for i in summary_raw.index]

    # 准备输出的长表格数据列表
    records = []

    # 生成长格式数据：每个区间 + 每个指标名 + 指标的均值
    for interval in summary_raw.index:
        for col in summary_raw.columns.levels[0]:
            # 过滤掉校区建筑面积的mean，若需要也可以保留
            if col == split_col:
                continue
            mean_val = summary_raw.loc[interval, (col, "mean")]
            records.append({
                split_col+"划分区间": interval,
                "目标列": col,
                "2021年": round(mean_val, 2)
            })

    # 转成DataFrame
    summary_df = pd.DataFrame(records)

    # 重置列顺序（防止乱序）
    summary_df = summary_df[[split_col+"划分区间", "目标列", "2021年"]]

    return df, boundaries, summary_df, summary_raw

def plot_scatter(df, split_col, target_col, thresholds, img_path):
    def assign_segment(x, thresholds, max_idx, min_idx):
        for i, t in enumerate(thresholds):
            if x <= t:
                if i == 0:
                    return f"分段1：[{min_idx},{t}]"
                else:
                    return f"分段{i+1}：({thresholds[i-1]},{t}]"
        return f"分段{len(thresholds)+1}：({thresholds[-1]},{max_idx}]"

    min_idx = round(df[split_col].min(), 4)
    max_idx = round(df[split_col].max(), 4)
    df["分段"] = df[split_col].apply(lambda x: assign_segment(x, thresholds, max_idx, min_idx))

    # 构建颜色映射
    unique_segments = sorted(df["分段"].unique())
    colors = plt.cm.get_cmap("Set2", len(unique_segments))
    color_map = {segment: colors(i) for i, segment in enumerate(unique_segments)}

    plt.figure(figsize=(7, 4))
    sns.set(font="AR PL UKai CN")

    for segment in unique_segments:
        segment_df = df[df["分段"] == segment]
        plt.scatter(
            segment_df[split_col],
            segment_df[target_col],
            label=segment,
            color=color_map[segment],
            alpha=0.8
        )

    plt.xlabel(split_col)
    plt.ylabel(target_col)
    plt.title(f"{split_col}与{target_col}的分段散点图")
    plt.legend()
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(img_path)
    plt.close()


@CatchException
def 区间划分(main_input, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    chatbot.append(["启动区间划分分析任务", "读取文件并自动处理sheet中..."])
    yield from update_ui(chatbot=chatbot, history=history)

    if os.path.isdir(main_input):
        excel_files = glob.glob(os.path.join(main_input, "*.xlsx"))
        if not excel_files:
            chatbot.append(["错误", f"目录 {main_input} 中未找到Excel文件"])
            yield from update_ui(chatbot=chatbot, history=history)
            return
        filepath = excel_files[0]
    else:
        filepath = main_input

    max_segments = 2  # 最大分段数
    min_ratio = 0.3  # 最小分段比例
    target_cols = ["单位面积成本", "人均单位成本", "平均服务面积", "保洁成本（2021）"]
    split_col = "校区建筑面积（平方米）"
    report_blocks = []
    final_report = ""
    image_map = {}

    df = pd.read_excel(filepath)
    chatbot.append(["读取文件", filepath])
    df, outlier_dict = clean_and_prepare(df, split_col, target_cols)
    for col, vals in outlier_dict.items():
        if vals.empty:
            continue
        chatbot.append([f"\n【{col}】离群值：", str(vals.values.tolist())])
    
    segmented_df, breakpoints, summary ,summary_ord= segment_by_area(
        df,
        split_col=split_col,
        target_cols=target_cols,
        method="tree",  # 可换成"maxdiff", "tree", "jenks"
        max_segments=max_segments,
        min_ratio=min_ratio
    )
    
    chatbot.append(["推荐分段边界：", str([float(x) for x in breakpoints])])
    chatbot.append(["分段统计：\n", summary.to_string()])
    for target_col in target_cols:
        img_path = f"scatter_{target_col.replace('/', '_')}.png"
        plot_scatter(segmented_df, split_col, target_col, breakpoints, img_path)
        image_map[f"scatter_{target_col.replace('/', '_')}"] = img_path
    summary_prompt = f"""
你是一位擅长数据分析的财务专家，请基于以下输入数据，总结**校区建筑面积区间划分下的成本趋势**，并用简洁专业的语言撰写一段分析性结论。请注意：

1. 根据输入数据中的区间划分（breakpoints），说明各区间的划分范围并在各指标上的平均值；
2. 指出各区间在单位面积成本、人均单位成本、平均服务面积、保洁总成本等方面的异同；并识别可能的规模效益（如面积大 → 单位成本低等）；
3. 输出应为一段文字分析，适合用于分析报告。

输入数据如下：
【区间划分】：{breakpoints}
【统计摘要】：\n{summary_ord.to_string()}
"""
    gpt_summary_reply = yield from request_gpt_model_in_new_thread_with_ui_alive(
        inputs=summary_prompt,
        inputs_show_user="区间划分后的成本趋势总结",
        llm_kwargs=llm_kwargs,
        chatbot=chatbot,
        history=history,
        sys_prompt="你是财务数据分析专家，擅长成本趋势和规模效益分析。"
    )

    summary_block = f"## 区间成本趋势分析总结\n{gpt_summary_reply.strip()}"
    final_report += "\n\n" + summary_block
    report_blocks.append(summary_block)

    chatbot.append(["所有可分析Sheet处理完成", final_report])
    history.append(("分析报告", final_report))
    res = write_history_to_file(history)
    promote_file_to_downloadzone(res, chatbot=chatbot)

    word_path = res.replace(".txt", ".docx").replace(".md", ".docx")
    save_markdown_to_docx(final_report, word_path, image_map=image_map, summary=summary)
    promote_file_to_downloadzone(word_path, chatbot=chatbot)

    for img in image_map.values():
        if os.path.exists(img):
            os.remove(img)

    chatbot.append(["已生成Word报告", word_path])
    chatbot.append(["已生成下载文件", res])
    yield from update_ui(chatbot=chatbot, history=history)
