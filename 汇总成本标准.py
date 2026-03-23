from toolbox import update_ui, CatchException, write_history_to_file, promote_file_to_downloadzone
import pandas as pd
import os, glob
import markdown2
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive

def save_markdown_to_styled_docx(md_text, path):
    html = markdown2.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    for elem in soup.descendants:
        if elem.name in ["h1", "h2", "p", "li", "table"]:
            text = elem.text.strip()
            if not text: continue
            if elem.name == "h1":
                p = doc.add_heading(text, level=1)
            elif elem.name == "h2":
                p = doc.add_heading(text, level=2)
            else:
                p = doc.add_paragraph(text)
            run = p.runs[0]
            run.font.name = '宋体'
            p.style.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
    doc.save(path)

@CatchException
def 保安岗位标准归纳插件(main_input, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    chatbot.append(["📥 插件启动", "正在读取 Excel 文件中的保安岗位数据..."])
    yield from update_ui(chatbot=chatbot, history=history)

    # 获取文件路径
    if os.path.isdir(main_input):
        files = glob.glob(os.path.join(main_input, "*.xlsx"))
        if not files:
            chatbot.append(["❌ 错误", "上传目录中未找到 Excel 文件"])
            yield from update_ui(chatbot=chatbot, history=history)
            return
        filepath = files[0]
    else:
        filepath = main_input

    # 读取保安sheet
    all_sheets = pd.read_excel(filepath, sheet_name=None)
    if "保安" not in all_sheets:
        chatbot.append(["⚠️ 警告", "未找到名为 '保安' 的 sheet"])
        yield from update_ui(chatbot=chatbot, history=history)
        return

    df = all_sheets["保安"]
    markdown_table = df.head(30).to_markdown(index=False)

    prompt = f"""
你是一位物业人力配置标准制定专家。客户上传了一份“保安岗位人员汇总表”（如下），请你根据表格内容，归纳并输出标准岗位配置表，要求如下：

1. 请你根据表格中的岗位名称（如“管理岗一班”、“客服岗二班”），归类出细化分类（例如：管理岗、监控岗、客服岗等）；
2. 每一类细化分类请设定一个岗位配置标准，例如：
   - “2岗/门”
   - “1岗/校区”
   - “3岗/区域”
   请结合班次、人数规模、命名方式推理出配置标准；
3. 请结合总成本与总人数字段，计算并输出人均成本（万元/人/年），数字保留1位小数；
4. 最终请严格输出如下 Markdown 表格格式：

| 物业服务模块 | 细化分类 | 岗位配置标准 | 人均成本标准（万元/人/年） |
|---------------|----------|----------------|-----------------------------|
| 保安           | 管理岗       | 1岗/校区         | 6.6                         |
| 保安           | 客服岗       | 2岗/值班岗       | 6.1                         |

❗请不要输出任何解释说明，只输出上述格式的表格。

表格原始数据如下：
{markdown_table}
"""

    reply = yield from request_gpt_model_in_new_thread_with_ui_alive(
        inputs=prompt,
        inputs_show_user="请生成岗位配置与成本标准表",
        llm_kwargs=llm_kwargs,
        chatbot=chatbot,
        history=history,
        sys_prompt="你是一位擅长归纳物业岗位数据的大模型专家"
    )

    history.append(("保安岗位配置表", reply.strip()))
    txt_path = write_history_to_file(history)
    promote_file_to_downloadzone(txt_path, chatbot=chatbot)

    docx_path = filepath.replace(".xlsx", "_保安岗位标准.docx")
    save_markdown_to_styled_docx(reply, docx_path)
    promote_file_to_downloadzone(docx_path, chatbot=chatbot)

    chatbot.append(["✅ 岗位标准表已生成", docx_path])
    yield from update_ui(chatbot=chatbot, history=history)
