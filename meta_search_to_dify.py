import os
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Optional
import json
import requests
from docx import Document
from loguru import logger

from toolbox import (
    CatchException,
    promote_file_to_downloadzone,
    report_exception,
    update_ui,
)
from shared_utils.config_loader import get_conf


def _load_settings(plugin_kwargs: Dict[str, Any]) -> Dict[str, Any]:
    """
    Read all runtime settings from env or plugin kwargs so users do not need to
    hardcode secrets inside the codebase.
    """
    def pick(name: str, default: Optional[str] = None) -> Optional[str]:
        return (
            plugin_kwargs.get(name)
            or os.getenv(name.upper())
            or default
        )

    return {
        "mita_api_key": pick("mita_api_key", "mk-E5C13B098F95036585319821025DAD74"),
        "mita_api_url": pick("mita_api_url", "https://metaso.cn/api/v1/search"),
        "mita_top_k": int(plugin_kwargs.get("top_k", 5)) if str(plugin_kwargs.get("top_k", "")).isdigit() else 5,
        "dify_api_key": pick("dify_api_key", "dataset-2OZXoojsUfMueuKd6afl6SIS"),
        "dify_dataset_id": pick("dify_dataset_id", "0c059f10-e1c4-4d91-a2cf-95ea50155882"),
        "dify_base_url": pick("dify_base_url", "http://localhost"),
    }

def _mita_search(query: str, settings: Dict[str, Any], proxies: Optional[dict]) -> List[Dict[str, Any]]:
    """
    调用秘塔搜索 API（POST JSON），返回网页结果列表。
    官方示例大致为：

    POST https://metaso.cn/api/v1/search
    {
        "q": "...",
        "scope": "webpage",
        "includeSummary": true,
        "size": "10",
        "includeRawContent": false,
        "conciseSnippet": true
    }
    """
    headers = {
        "Authorization": f"Bearer {settings['mita_api_key']}",
        "Accept": "application/json",
        "Content-Type": "application/json",
    }

    payload = {
        "q": query,
        "scope": "webpage",                         # 先固定网页搜索
        "includeSummary": True,                     # 需要摘要，方便写入 Word
        "size": str(settings["mita_top_k"]),        # 官方示例里是字符串，这里也转成字符串
        "includeRawContent": False,                 # 如需全文可以改成 True
        "conciseSnippet": True,
    }

    resp = requests.post(
        settings["mita_api_url"],
        json=payload,
        headers=headers,
        timeout=60,
        proxies=proxies,
    )
    resp.raise_for_status()
    data = resp.json()

    # 秘塔搜索实际返回格式通常是 {"webpages": [ {...}, {...} ]}
    if isinstance(data, dict):
        # 1) 直接有 webpages
        if "webpages" in data and isinstance(data["webpages"], list):
            return data["webpages"]

        # 2) 有些封装可能是 {"data": {"webpages": [...]}}
        if "data" in data and isinstance(data["data"], dict):
            inner = data["data"]
            if "webpages" in inner and isinstance(inner["webpages"], list):
                return inner["webpages"]

        # 3) 兜底：沿用之前的通用提取逻辑
        for key in ["results", "data", "items"]:
            if key in data and isinstance(data[key], list):
                return data[key]

    elif isinstance(data, list):
        return data

    return []


# def _mita_search(query: str, settings: Dict[str, Any], proxies: Optional[dict]) -> List[Dict[str, Any]]:
#     headers = {"Authorization": f"Bearer {settings['mita_api_key']}"}
#     params = {"q": query, "limit": settings["mita_top_k"]}
#     resp = requests.get(
#         settings["mita_api_url"],
#         params=params,
#         headers=headers,
#         timeout=60,
#         proxies=proxies,
#     )
#     resp.raise_for_status()
#     data = resp.json()
#     # 兼容不同接口返回格式，尽量尝试提取常见字段
#     if isinstance(data, dict):
#         for key in ["results", "data", "items"]:
#             if key in data and isinstance(data[key], list):
#                 return data[key]
#     if isinstance(data, list):
#         return data
#     return []


def _build_word_doc(query: str, results: List[Dict[str, Any]]) -> str:
    doc = Document()
    doc.add_heading(f"秘塔搜索结果：{query}", level=1)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"生成时间：{now}")

    if not results:
        doc.add_paragraph("未获取到搜索结果。")
    else:
        for idx, item in enumerate(results, 1):
            title = item.get("title") or item.get("name") or f"结果 {idx}"
            summary = item.get("summary") or item.get("abstract") or item.get("snippet") or ""
            url = item.get("url") or item.get("link") or ""
            published = item.get("published_at") or item.get("date") or item.get("time") or ""

            doc.add_heading(f"{idx}. {title}", level=2)
            if summary:
                doc.add_paragraph(summary)
            if url:
                doc.add_paragraph(f"URL: {url}")
            if published:
                doc.add_paragraph(f"时间: {published}")

    tmp_dir = tempfile.mkdtemp(prefix="mita_search_")
    filename = f"{query[:40].replace(' ', '_') or 'search'}.docx"
    output_path = os.path.join(tmp_dir, filename)
    doc.save(output_path)
    return output_path

def _get_dify_doc_form(settings: Dict[str, Any]) -> Optional[str]:
    """
    查询当前知识库的 doc_form（text_model / qa_model / hierarchical_model 等）。
    避免出现 “doc_form is different from the dataset doc_form” 的 400 错误。
    """
    base_url = settings["dify_base_url"].rstrip("/")
    dataset_id = settings["dify_dataset_id"]
    url = f"{base_url}/v1/datasets/{dataset_id}"
    headers = {"Authorization": f"Bearer {settings['dify_api_key']}"}

    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        data = resp.json()

        # 有的版本是 {"data": {...}}，有的是直接 {...}
        if isinstance(data, dict):
            dataset = data.get("data") or data
            doc_form = dataset.get("doc_form")
            logger.info(f"[Dify] dataset {dataset_id} doc_form = {doc_form}")
            return doc_form
    except Exception as e:
        logger.warning(f"[Dify] 获取 doc_form 失败：{e}")

    return None


def _upload_to_dify(doc_path: str, settings: Dict[str, Any]) -> Dict[str, Any]:
    """
    调用 Dify 文档接口：
    POST /v1/datasets/{dataset_id}/document/create-by-file

    通过上传文件在现有知识库中创建新文档。
    只传 doc_form，其余索引规则完全沿用知识库本身的配置，
    避免前端上传和 API 上传的行为不一致。
    """
    base_url = settings["dify_base_url"].rstrip("/")
    dataset_id = settings["dify_dataset_id"]
    upload_url = f"{base_url}/v1/datasets/{dataset_id}/document/create-by-file"

    # 1. 获取知识库的 doc_form（text_model / qa_model / hierarchical_model 等）
    doc_form = _get_dify_doc_form(settings)

    # 2. 只把 doc_form 传给 Dify，其余全部使用 dataset 默认 process_rule
    config: Dict[str, Any] = {}
    if doc_form:
        config["doc_form"] = doc_form

    headers = {
        "Authorization": f"Bearer {settings['dify_api_key']}",
    }

    # multipart/form-data：字段名 data + file
    data = {
        "data": json.dumps(config, ensure_ascii=False)
    }

    with open(doc_path, "rb") as f:
        files = {
            "file": (
                os.path.basename(doc_path),
                f,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        resp = requests.post(
            upload_url,
            headers=headers,
            data=data,
            files=files,
            timeout=180,
        )

    try:
        resp.raise_for_status()
    except requests.HTTPError as e:
        raise RuntimeError(
            f"Dify upload failed: {resp.status_code} {resp.text}"
        ) from e

    try:
        return resp.json()
    except Exception:
        return {"raw": resp.text}


# def _upload_to_dify(doc_path: str, settings: Dict[str, Any]) -> Dict[str, Any]:
#     upload_url = f"{settings['dify_base_url'].rstrip('/')}/v1/datasets/{settings['dify_dataset_id']}/documents/upload"
#     headers = {"Authorization": f"Bearer {settings['dify_api_key']}"}
#     files = {
#         "file": (
#             os.path.basename(doc_path),
#             open(doc_path, "rb"),
#             "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         )
#     }
#     resp = requests.post(upload_url, headers=headers, files=files, timeout=180)
#     resp.raise_for_status()
#     try:
#         return resp.json()
#     except Exception:
#         return {"raw": resp.text}


@CatchException
def 秘塔搜索入库(txt, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    """
    调用秘塔搜索，生成 Word 文档并上传到 Dify 知识库。
    """
    chatbot.append([
        "秘塔搜索入库",
        "调用秘塔搜索 → 生成 Word → 上传到 Dify 知识库",
    ])
    yield from update_ui(chatbot=chatbot, history=history)

    query = (txt or "").strip()
    if not query:
        chatbot[-1] = [chatbot[-1][0], "请输入搜索关键词后再试一次。"]
        yield from update_ui(chatbot=chatbot, history=history)
        return

    settings = _load_settings(plugin_kwargs or {})
    missing = [k for k in ["mita_api_key", "dify_api_key", "dify_dataset_id"] if not settings.get(k)]
    if missing:
        report_exception(
            chatbot,
            history,
            a="秘塔搜索入库",
            b=f"缺少配置：{', '.join(missing)}。请在环境变量或插件高级参数中提供它们。",
        )
        yield from update_ui(chatbot=chatbot, history=history)
        return

    proxies = None
    try:
        proxies = get_conf("proxies")
    except Exception:
        proxies = None

    try:
        chatbot[-1] = [chatbot[-1][0], "正在调用秘塔搜索..."]
        yield from update_ui(chatbot=chatbot, history=history)
        results = _mita_search(query, settings, proxies)
    except Exception as e:
        logger.error(e)
        report_exception(
            chatbot,
            history,
            a="秘塔搜索失败",
            b=str(e),
        )
        yield from update_ui(chatbot=chatbot, history=history)
        return

    chatbot[-1] = [chatbot[-1][0], f"搜索完成，生成 Word 文档中（{len(results)} 条结果）..."]
    yield from update_ui(chatbot=chatbot, history=history)

    doc_path = _build_word_doc(query, results)
    promote_file_to_downloadzone(doc_path, chatbot=chatbot)

    try:
        chatbot[-1] = [chatbot[-1][0], "上传到 Dify 知识库中..."]
        yield from update_ui(chatbot=chatbot, history=history)
        upload_result = _upload_to_dify(doc_path, settings)
    except Exception as e:
        logger.error(e)
        report_exception(
            chatbot,
            history,
            a="上传到 Dify 失败",
            b=str(e),
        )
        yield from update_ui(chatbot=chatbot, history=history)
        return

    chatbot[-1] = [
        chatbot[-1][0],
        f"完成：已上传到 Dify 知识库。\n\n响应：{upload_result}",
    ]
    yield from update_ui(chatbot=chatbot, history=history)

