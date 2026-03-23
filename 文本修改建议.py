from toolbox import update_ui
from toolbox import CatchException, report_exception
from toolbox import write_history_to_file, promote_file_to_downloadzone
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive

fast_debug = False


def 解析word文本建议(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt):
    """
    批量解析 Word 文档并对内容生成“文本修改建议”与“修订稿”两部分输出，汇总到一个 Word 报告。
    - file_manifest: 需要处理的文件清单（绝对或相对路径）
    - project_folder: 项目根目录（用于相对路径展示）
    - llm_kwargs: 大模型参数（包含 llm_model 等）
    - plugin_kwargs: 插件参数（可选：style、domain、audience 等提示）
    - chatbot/history/system_prompt: GPT_Academic 运行时上下文
    """
    import os
    from docx import Document  # 用于生成报告
    from datetime import datetime

    # 读取用户的可选风格/领域偏好
    style_hint = plugin_kwargs.get("style", "").strip() if isinstance(plugin_kwargs, dict) else ""
    domain_hint = plugin_kwargs.get("domain", "").strip() if isinstance(plugin_kwargs, dict) else ""
    audience_hint = plugin_kwargs.get("audience", "").strip() if isinstance(plugin_kwargs, dict) else ""

    # 新建报告文档
    report = Document()
    report.add_heading("文本修改建议报告", level=0)
    report.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if style_hint or domain_hint or audience_hint:
        report.add_paragraph("风格/领域/读者偏好：")
        if style_hint: report.add_paragraph(f"- 目标风格：{style_hint}")
        if domain_hint: report.add_paragraph(f"- 领域背景：{domain_hint}")
        if audience_hint: report.add_paragraph(f"- 预期读者：{audience_hint}")

    for index, fp in enumerate(file_manifest):
        # 读取文档内容（.docx 优先；.doc 走 win32com）
        if fp.split(".")[-1].lower() == "docx":
            from docx import Document as DocReader
            doc = DocReader(fp)

            # 提取文字
            text_content = "\n".join([para.text for para in doc.paragraphs])

            # 提取表格为制表符分隔行
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_content.append("\t".join(row_text))

            file_content = text_content
            if table_content:
                file_content += "\n表格内容（以制表符分隔）：\n" + "\n".join(table_content)
        else:
            try:
                import os as _os
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.visible = False
                doc = word.Documents.Open(_os.getcwd() + '/' + fp)
                doc = word.ActiveDocument
                file_content = doc.Range().Text
                doc.Close()
                word.Quit()
            except Exception as e:
                raise RuntimeError('请先将 .doc 文档转换为 .docx 文档，或在 Windows 环境下安装 pywin32 后重试。错误：%s' % e)

        # 将文本分片以满足 token 限制
        from crazy_functions.pdf_fns.breakdown_txt import breakdown_text_to_satisfy_token_limit
        from request_llms.bridge_all import model_info
        max_token = model_info[llm_kwargs['llm_model']]['max_token']
        TOKEN_LIMIT_PER_FRAGMENT = max_token * 3 // 4
        fragments = breakdown_text_to_satisfy_token_limit(
            txt=file_content,
            limit=TOKEN_LIMIT_PER_FRAGMENT,
            llm_model=llm_kwargs['llm_model']
        )

        this_doc_history = []
        import os as _os
        report.add_heading(f"文件：{_os.path.basename(fp)}", level=1)

        # 针对每个片段生成“修改建议 + 修订稿”
        for i, frag in enumerate(fragments):
            # 构造提示
            preferences = []
            if style_hint: preferences.append(f"目标风格：{style_hint}")
            if domain_hint: preferences.append(f"领域背景：{domain_hint}")
            if audience_hint: preferences.append(f"预期读者：{audience_hint}")
            pref_text = "；".join(preferences) if preferences else "（无特别偏好）"

            i_say = (
                "你是一名严谨的学术与技术写作编辑，现请对以下文档片段给出“文本修改建议”，并提供修订后的版本。\n"
                "请从以下维度提出**可操作**的修改建议：\n"
                "1) 清晰性与逻辑性（是否存在歧义、跳跃、堆砌）；\n"
                "2) 结构与衔接（段落组织、过渡、主题句）；\n"
                "3) 简练与去冗（长句拆分、口语化表达收敛）；\n"
                "4) 术语一致与格式规范（学术/工程名词、标点与单位、大小写/中英文混排）；\n"
                "5) 语法与用词（时态、主谓一致、搭配、拼写）；\n"
                "6) 语域与读者（面向对象匹配度、正式度、客观性）。\n\n"
                f"编辑偏好：{pref_text}\n\n"
                "请以**Markdown**结构化输出，严格包含以下三部分：\n"
                "【A. 修改要点清单】（列表列出关键问题与对应改进策略，尽量引用原句片段）\n"
                "【B. 逐段修订稿】（在保持含义不变的前提下，给出优化后的完整文本；保留段落结构）\n"
                "【C. 术语与风格统一建议】（罗列需在全文保持一致的术语、单位、符号与表达规则）\n\n"
                f"文件名：{_os.path.relpath(fp, project_folder)}，第 {i+1}/{len(fragments)} 片段内容如下：\n"
                f"```text\n{frag}\n```"
            )

            i_say_show_user = f'对 { _os.path.abspath(fp) } 的第 {i+1}/{len(fragments)} 个片段生成文本修改建议。'

            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say_show_user,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=[],
                sys_prompt="文本修改建议。"
            )

            # 刷新前端
            chatbot[-1] = (i_say_show_user, gpt_say)
            history.extend([i_say_show_user, gpt_say])
            this_doc_history.extend([i_say_show_user, gpt_say])

            # 写入报告
            report.add_heading(f"片段 {i+1} 修改建议与修订稿", level=2)
            report.add_paragraph(gpt_say)

        # 整体风格与术语汇总（仅当存在多个片段）
        if len(fragments) > 1:
            i_say = (
                "基于前述各片段的编辑结果，请为该文档生成**全局写作风格与术语统一指南**，"
                "并输出一个不超过 200 字的“全文修订摘要”（突出主要问题与总体改进方向）。"
            )
            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=this_doc_history,
                sys_prompt="文本修改建议。"
            )
            history.extend([i_say, gpt_say])
            this_doc_history.extend([i_say, gpt_say])

            # 写入报告
            report.add_heading("全局写作风格与术语统一指南 + 全文修订摘要", level=2)
            report.add_paragraph(gpt_say)

    # 保存报告
    import os as _os
    report_path = _os.path.join(project_folder, "文本修改建议报告.docx")
    report.save(report_path)

    # 同时保存 txt 历史记录并推送下载
    res = write_history_to_file(history)
    promote_file_to_downloadzone(res, chatbot=chatbot)
    promote_file_to_downloadzone(report_path, chatbot=chatbot)

    chatbot.append(("最终报告已生成", f"报告文件路径：{report_path}"))
    yield from update_ui(chatbot=chatbot, history=history)


@CatchException
def 文本修改建议(txt, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    """
    插件功能说明：
    - 批量读取 Word 文档（.docx；Windows 环境下可选 .doc），抽取正文与表格；
    - 调用大模型对文本进行学术/工程向“文本修改建议”与“修订稿”生成；
    - 自动汇总并导出为《文本修改建议报告.docx》，并在对话区提供下载。

    使用说明：
    - txt 可为单一 .docx/.doc 文件路径，也可以是包含若干 Word 文件的目录路径；
    - 可在 plugin_kwargs 中传入 style/domain/audience 等偏好字段用于定制编辑风格。
    """
    import glob, os

    # 前端提示
    chatbot.append([
        "函数插件功能？",
        "批量读取Word文档，抽取正文与表格，调用大模型生成“文本修改建议 + 修订稿”，并导出Word报告。"
    ])
    yield from update_ui(chatbot=chatbot, history=history)

    # 依赖检查
    try:
        from docx import Document  # noqa: F401
    except Exception:
        report_exception(
            chatbot, history,
            a=f"解析项目: {txt}",
            b="缺少依赖，请安装：```pip install python-docx pywin32```。"
        )
        yield from update_ui(chatbot=chatbot, history=history)
        return

    history = []

    # 路径校验
    if os.path.exists(txt):
        project_folder = txt
    else:
        if txt == "":
            txt = '空输入'
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"找不到路径或无权限访问: {txt}")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    # 生成文件清单
    if txt.lower().endswith('.docx') or txt.lower().endswith('.doc'):
        file_manifest = [txt]
        project_folder = os.path.dirname(os.path.abspath(txt)) or os.getcwd()
    else:
        project_folder = os.path.abspath(project_folder)
        file_manifest = [f for f in glob.glob(f'{project_folder}/**/*.docx', recursive=True)]
        # 如需支持 .doc，可解除下一行注释（需 Windows + pywin32）：
        # file_manifest += [f for f in glob.glob(f'{project_folder}/**/*.doc', recursive=True)]

    if len(file_manifest) == 0:
        report_exception(chatbot, history, a=f"解析项目: {txt}", b="未找到任何 Word 文件。")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    # 进入主流程
    yield from 解析word文本建议(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt)
