from toolbox import update_ui, CatchException, report_exception, write_history_to_file, promote_file_to_downloadzone
import pandas as pd
import numpy as np
from sklearn.linear_model import LinearRegression
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive
import os
import glob
import matplotlib.pyplot as plt
import seaborn as sns
import markdown2
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import matplotlib
import re
from sklearn.preprocessing import StandardScaler
from collections import defaultdict, Counter
import textwrap
import statsmodels.api as sm

matplotlib.rcParams['font.family'] = 'AR PL UKai CN'
matplotlib.rcParams['font.sans-serif'] = ['AR PL UKai CN']
matplotlib.use('Agg')
matplotlib.rcParams['axes.unicode_minus'] = False 

def plot_feature_importance_bar(values, features, title, save_path):
    plt.figure(figsize=(7, 4))
    plt.rcParams['font.family'] = 'AR PL UKai CN'
    sns.set(font="AR PL UKai CN")  
    clean_features = [str(f).replace("−", "-") for f in features]
    clean_values = [float(str(v).replace("−", "-")) for v in values]
    clean_features_wrapped = [textwrap.fill(label, 5) for label in clean_features]

    ax = sns.barplot(x=clean_features_wrapped, y=clean_values)
    for i, v in enumerate(clean_values):
        ax.text(v + 0.01, i, f"{v:.2f}", va='center')
    plt.title(title)
    plt.xlabel('各动因的重要性')
    plt.tight_layout()
    plt.savefig(save_path, dpi=150)
    plt.close()

def plot_top_feature_vs_target_scatter(df, target_col, top_feature, title, save_path):
    # 清洗数据：去除NaN和0
    filtered_df = df[[target_col, top_feature]].dropna()
    filtered_df = filtered_df[(filtered_df[target_col] != 0) & (filtered_df[top_feature] != 0)]

    if filtered_df.empty:
        print(f"跳过绘图：{title}，因无有效数据")
        return

    plt.figure(figsize=(7, 4))
    sns.set(font="AR PL UKai CN")
    sns.scatterplot(x=filtered_df[target_col], y=filtered_df[top_feature], color='dodgerblue', s=80)
    lowess = sm.nonparametric.lowess
    z = lowess(filtered_df[top_feature], filtered_df[target_col], frac=0.3)  # frac 控制平滑程度
    plt.plot(z[:, 0], z[:, 1], color='green', linestyle='-', label='LOWESS 拟合')

    plt.xlabel(target_col)
    plt.ylabel(top_feature)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()

def save_markdown_to_docx(md_text, path, image_map=None):
    html = markdown2.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    img_prefixes = ["bar_", "scatter_"]  
    current_section_key = None  # 用于匹配对应图像

    for elem in soup.descendants:
        if elem.name in ["h1", "h2", "p", "li"]:
            text = elem.text.strip()
            if not text:
                continue

            if elem.name == "h1":
                p = doc.add_heading(text, level=1)
            elif elem.name == "h2":
                p = doc.add_heading(text, level=2)
                # 尝试从标题中解析 sheet 和 y_col 名称
                match = re.search(r"工作表：(.*?) - 变量：(.*)", text)
                if match:
                    sheet, y_col = match.groups()
                    current_section_key = f"{sheet}_{y_col}"
            elif elem.name == "li":
                p = doc.add_paragraph(text, style="List Bullet")
            else:
                p = doc.add_paragraph(text)

            run = p.runs[0]
            run.font.name = 'AR PL UKai CN'
            p.style.font.name = 'AR PL UKai CN'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'AR PL UKai CN')
 
            # 插入图片在标题之后
            if current_section_key and image_map:
                for prefix in img_prefixes:
                    img_key = prefix + current_section_key
                    img_path = image_map.get(img_key)
                    if img_path and os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(5.5))
                current_section_key = None  # 插图完成，清除 key
    doc.save(path)

@CatchException
def 成本动因分析_LLM多Sheet版(main_input, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    chatbot.append(["启动成本动因分析任务", "读取文件并自动处理多个sheet..."])
    yield from update_ui(chatbot=chatbot, history=history)

    if os.path.isdir(main_input):
        excel_files = glob.glob(os.path.join(main_input, "*.xlsx"))
        if not excel_files:
            chatbot.append(["错误", f"目录 {main_input} 中未找到Excel文件"])
            yield from update_ui(chatbot=chatbot, history=history)
            return
        filepath = excel_files[0]
    else:
        filepath = main_input
   
    all_sheets = pd.read_excel(filepath, sheet_name=None)
    report_blocks = []
    image_map = {}
    all_model_results = []

    for sheet_name, df in all_sheets.items():
        df.columns = [str(col).strip() for col in df.columns]
        df.replace("–", np.nan, inplace=True)
        df = df.applymap(lambda x: str(x).replace(",", "") if isinstance(x, str) else x)
        df = df.apply(pd.to_numeric, errors="ignore")

        # valid_cols = []
        # for col in df.columns:
        #     try:
        #         pd.to_numeric(subdf[col], errors='raise')
        #         valid_cols.append(col)
        #     except:
        #         continue

        target_keywords = ["金额", "费用", "支出", "成本", "cost", "价格", "额"]
        y_cols = [col for col in df.columns if any(re.search(k, col, re.IGNORECASE) for k in target_keywords)]
        x_cols  = [col for col in df.columns if col not in y_cols and col != '序号' and col != '学校名称' and col != '校区']
        chatbot.append([f"处理工作表 {sheet_name}", f"找到因变量：{', '.join(y_cols)}，自变量：{', '.join(x_cols)}"])
        if not y_cols or not x_cols:
            chatbot.append([f"跳过工作表 {sheet_name}", "未找到有效的因变量或自变量"])
            continue

        for y_col in y_cols:
            subdf = df[[y_col] + x_cols].dropna()
            if len(subdf) < len(x_cols) + 3:
                continue

            y = subdf[y_col].astype(float).values
            X = subdf[x_cols].astype(float).values
            scaler = StandardScaler()
            X_scaled = scaler.fit_transform(X)
            model = LinearRegression().fit(X_scaled, y)

            r2 = model.score(X_scaled, y)
            coef = model.coef_
            intercept = model.intercept_
            std_unit_cost = (y / (X.sum(axis=1) + 1e-5)).std()

            eq_terms = [f"{coef[i]:.2f}×{x_cols[i]}" for i in range(len(x_cols))]
            # equation = f"{y_col} = {' + '.join(eq_terms)} + {intercept:.2f}"
            image_path1 = f"tmp_importance_{sheet_name}_{y_col}.png"
            plot_feature_importance_bar(coef, x_cols, f"{y_col}的回归系数", image_path1)
            image_map[f"bar_{sheet_name}_{y_col}"] = image_path1

            max_idx = np.argmax(np.abs(coef))
            most_important_feature = x_cols[max_idx]
            image_path2 = f"tmp_scatter_{sheet_name}_{y_col}_{most_important_feature}.png"
            plot_top_feature_vs_target_scatter(df, y_col, most_important_feature,
                                                f"{y_col}与{most_important_feature}", image_path2)
            image_map[f"scatter_{sheet_name}_{y_col}"] = image_path2

            all_model_results.append({
                "sheet": sheet_name,
                "target": y_col,
                "features": x_cols,
                "coef": model.coef_.tolist(),
                "score": r2
            })
            summary_prompt = f"""
你是一个成本动因分析专家，请分析以下回归结果并撰写一段分析报告，内容需包括：

1. 对回归结果的解释（例如哪个自变量影响最大）
2. R² 表示拟合程度，是否足够好？
3. 给出基于上述分析的管理建议或洞察

数据如下：
工作表：{sheet_name}
因变量：{y_col}
R² = {r2:.3f}
单位成本波动 std = {std_unit_cost:.2f}
回归系数如下：
{chr(10).join([f"{x_cols[i]}: {coef[i]:.2f}" for i in range(len(x_cols))])}
请以正式风格撰写分析段落。
"""
            gpt_reply = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=summary_prompt,
                inputs_show_user=f"Sheet: {sheet_name} - 因变量: {y_col}",
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=history,
                sys_prompt="你是成本动因分析专家，请撰写正式分析报告段落。"
            )

            block = f"## 工作表：{sheet_name} - 变量：{y_col}\n{gpt_reply.strip()}"
            report_blocks.append(block)
            
    results_by_target = defaultdict(lambda: defaultdict(dict))
    chatbot.append(["所有工作表分析完成", "正在整理结果..."])
    for result in all_model_results:
        target = result["target"]
        year = result["sheet"]
        results_by_target[target][year] = {
            "coef": np.array(result["coef"]),
            "score": result["score"],
            "features": result["features"]
        }

    target_summary = []

    for target, year_data in results_by_target.items():
        target_report = f"【因变量：{target}】\n"

        top_features_by_year = {}
        all_features = []
        for year, info in year_data.items():
            coef_abs = np.abs(info["coef"])
            top_index = np.argmax(coef_abs)
            top_feat = info["features"][top_index]
            top_features_by_year[year] = top_feat
            all_features.append(top_feat)

        counter = Counter(all_features)
        common_factors = [k for k, v in counter.items() if v >= 2]

        target_report += f"- 各年份第一成本动因：{top_features_by_year}\n"
        if common_factors:
            target_report += f"- 多年反复出现的重要因子：{common_factors}\n"
        else:
            target_report += f"- 各年份主导因子差异较大，缺乏共性\n"

        scores = {year: round(info['score'], 3) for year, info in year_data.items()}
        target_report += f"- 拟合效果（R²）：{scores}\n"

        target_summary.append(target_report)
    cross_year_summary = "\n\n".join(target_summary)
    final_report = "# 成本动因分析报告\n\n" + "\n\n".join(report_blocks)
    summary_of_summaries = "\n\n".join(report_blocks)

    summary_prompt1 = f"""
你是一位经验丰富的财务分析专家，如果每个工作表中有相同的因变量，请进行比较他们的各自的成本动因并进行总结，生成成本动因分析报告，撰写一个整体汇总段落，内容包括：

1. 先说明哪个自变量的回归系数最大；
2. 然后汇总各表中相同的因变量由各自成本动因的共性与差异；
3. 哪些因子在多个年份（多个sheet表）中反复表现出较高的重要性;
4. 总结和综合建议

请以正式风格撰写该汇总段落，不要逐年罗列，请进行总结与概括。
这是汇总的数据：{cross_year_summary}
"""

    gpt_summary_reply1 = yield from request_gpt_model_in_new_thread_with_ui_alive(
        inputs=summary_prompt1,
        inputs_show_user="最终汇总总结",
        llm_kwargs=llm_kwargs,
        chatbot=chatbot,
        history=history,
        sys_prompt="你是财务分析专家，请撰写该文件的成本动因分析的总结性段落。"
    )

    summary_block = f"## 全年成本动因分析总结\n{gpt_summary_reply1.strip()}"
    final_report += "\n\n" + summary_block
    report_blocks.append(summary_block)

    chatbot.append(["所有可分析Sheet处理完成", final_report])
    history.append(("分析报告", final_report))
    res = write_history_to_file(history, file_basename="成本动因分析报告.md")
    promote_file_to_downloadzone(res, chatbot=chatbot)

    word_path = res.replace(".txt", ".docx").replace(".md", ".docx")
    save_markdown_to_docx(final_report, word_path, image_map=image_map)
    promote_file_to_downloadzone(word_path, chatbot=chatbot)

    for img in image_map.values():
        if os.path.exists(img):
            os.remove(img)

    chatbot.append(["已生成Word报告", word_path])
    chatbot.append(["已生成下载文件", res])
    yield from update_ui(chatbot=chatbot, history=history)