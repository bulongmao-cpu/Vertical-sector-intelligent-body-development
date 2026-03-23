from toolbox import update_ui
from toolbox import CatchException, report_exception
from toolbox import write_history_to_file, promote_file_to_downloadzone
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive

fast_debug = False


def 解析word趋势(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt):
    import os
    from docx import Document  # 用于生成报告
    from datetime import datetime

    # 新建报告文档
    report = Document()
    report.add_heading("数据趋势分析报告", level=0)
    report.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    for index, fp in enumerate(file_manifest):
        if fp.split(".")[-1] == "docx":
            from docx import Document as DocReader
            doc = DocReader(fp)

            # 提取文字
            text_content = "\n".join([para.text for para in doc.paragraphs])

            # 提取表格
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_content.append("\t".join(row_text))

            file_content = text_content + "\n表格内容：\n" + "\n".join(table_content)
        else:
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.visible = False
                doc = word.Documents.Open(os.getcwd() + '/' + fp)
                doc = word.ActiveDocument
                file_content = doc.Range().Text
                doc.Close()
                word.Quit()
            except:
                raise RuntimeError('请先将.doc文档转换为.docx文档。')

        # 分片
        from crazy_functions.pdf_fns.breakdown_txt import breakdown_text_to_satisfy_token_limit
        from request_llms.bridge_all import model_info
        max_token = model_info[llm_kwargs['llm_model']]['max_token']
        TOKEN_LIMIT_PER_FRAGMENT = max_token * 3 // 4
        fragments = breakdown_text_to_satisfy_token_limit(
            txt=file_content,
            limit=TOKEN_LIMIT_PER_FRAGMENT,
            llm_model=llm_kwargs['llm_model']
        )

        this_paper_history = []
        report.add_heading(f"文件：{os.path.basename(fp)}", level=1)

        for i, frag in enumerate(fragments):
            i_say = f'请对下面的文档片段进行**数据趋势分析**，识别变化规律、显著波动和潜在原因。文件名：{os.path.relpath(fp, project_folder)}，内容如下：```{frag}```'
            i_say_show_user = f'对 {os.path.abspath(fp)} 的第 {i+1}/{len(fragments)} 个片段进行趋势分析。'

            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say_show_user,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=[],
                sys_prompt="数据趋势分析。"
            )

            chatbot[-1] = (i_say_show_user, gpt_say)
            history.extend([i_say_show_user, gpt_say])
            this_paper_history.extend([i_say_show_user, gpt_say])

            # 写入报告
            report.add_heading(f"片段 {i+1} 分析结果", level=2)
            report.add_paragraph(gpt_say)

        # 整体趋势总结
        if len(fragments) > 1:
            i_say = f"根据以上分析结果，总结整个文档 {os.path.abspath(fp)} 的总体趋势与结论。"
            gpt_say = yield from request_gpt_model_in_new_thread_with_ui_alive(
                inputs=i_say,
                inputs_show_user=i_say,
                llm_kwargs=llm_kwargs,
                chatbot=chatbot,
                history=this_paper_history,
                sys_prompt="数据趋势分析。"
            )
            history.extend([i_say, gpt_say])
            this_paper_history.extend([i_say, gpt_say])

            # 写入报告
            report.add_heading("整体趋势总结", level=2)
            report.add_paragraph(gpt_say)

    # 保存报告
    report_path = os.path.join(project_folder, "数据趋势分析报告.docx")
    report.save(report_path)

    # 同时保存 txt 历史记录
    res = write_history_to_file(history)
    promote_file_to_downloadzone(res, chatbot=chatbot)
    promote_file_to_downloadzone(report_path, chatbot=chatbot)

    chatbot.append(("最终报告已生成", f"报告文件路径：{report_path}"))
    yield from update_ui(chatbot=chatbot, history=history)


@CatchException
def 数据趋势分析(txt, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    import glob, os

    chatbot.append([
        "函数插件功能？",
        "批量读取Word文档，提取文字与表格，调用大模型进行趋势分析，并导出Word报告。"])
    yield from update_ui(chatbot=chatbot, history=history)

    # 依赖检查
    try:
        from docx import Document
    except:
        report_exception(chatbot, history,
                         a=f"解析项目: {txt}",
                         b=f"缺少依赖，请安装：```pip install python-docx pywin32```。")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    history = []

    if os.path.exists(txt):
        project_folder = txt
    else:
        if txt == "": txt = '空输入'
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"找不到路径或无权限访问: {txt}")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    if txt.endswith('.docx') or txt.endswith('.doc'):
        file_manifest = [txt]
    else:
        file_manifest = [f for f in glob.glob(f'{project_folder}/**/*.docx', recursive=True)]
        # file_manifest = [f for f in glob.glob(f'{project_folder}/**/*.docx', recursive=True)] + \
        #                 [f for f in glob.glob(f'{project_folder}/**/*.doc', recursive=True)]

    if len(file_manifest) == 0:
        report_exception(chatbot, history, a=f"解析项目: {txt}", b=f"未找到任何Word文件。")
        yield from update_ui(chatbot=chatbot, history=history)
        return

    yield from 解析word趋势(file_manifest, project_folder, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt)
