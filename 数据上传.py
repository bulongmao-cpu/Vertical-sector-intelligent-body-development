from openpyxl import load_workbook
from collections import defaultdict
import json
from openpyxl.styles import PatternFill
from sqlalchemy import create_engine, text
from sqlalchemy import text
import math
from sklearn.tree import DecisionTreeRegressor
from scipy.stats import f_oneway
import mapclassify
from toolbox import update_ui, CatchException, report_exception, write_history_to_file, promote_file_to_downloadzone
import pandas as pd
import numpy as np
from sklearn.linear_model import LinearRegression
from crazy_functions.crazy_utils import request_gpt_model_in_new_thread_with_ui_alive
import os
import glob
import matplotlib.pyplot as plt
import seaborn as sns
import markdown2
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import matplotlib
import re
from sklearn.preprocessing import StandardScaler
from collections import defaultdict, Counter
from matplotlib.font_manager import FontProperties
from crazy_functions.crazy_utils import request_input_from_user

class MySQLWriter:
    def __init__(self, database: str):
        """
        初始化 MySQLWriter
        """
        self.database = database
        self.engine = create_engine(
            f"mysql+pymysql://root:123456@localhost:3306/{database}?charset=utf8mb4"
        )

    # ==================== 增（Insert） ====================
    def insert_df(self, df: pd.DataFrame, table_name: str, if_exists: str = "replace"):
        df.to_sql(name=table_name, con=self.engine, if_exists=if_exists, index=False)
        print(f"插入 {len(df)} 行数据到 {table_name} （模式：{if_exists}）")

    def insert_one(self, table_name: str, data: dict):
        cols = ", ".join(data.keys())
        placeholders = ", ".join([f":{k}" for k in data.keys()])
        sql = f"INSERT INTO {table_name} ({cols}) VALUES ({placeholders})"
        with self.engine.connect() as conn:
            conn.execute(text(sql), data)
            conn.commit()
        print(f"插入单条数据到 {table_name}")

    # ==================== 查（Select） ====================
    def select_df(self, table_name: str, columns: list[str] = None, where: str = None) -> pd.DataFrame:
        def quote_identifier(name):
            return f"`{name}`"

        if columns is None:
            col_str = "*"
        else:
            col_str = ", ".join(quote_identifier(c) for c in columns)

        table_name = quote_identifier(table_name)   # 处理包含中文/空格/特殊字符的表名
        sql = f"SELECT {col_str} FROM {table_name}"

        if where:
            sql += f" WHERE {where}"

        df = pd.read_sql(sql, self.engine)
        print(f"从 {table_name} 读取 {len(df)} 行数据")
        return df
    
    # ==================== 改（Update） ====================
    def update(self, table_name: str, updates: dict, where: str):
        set_clause = ", ".join([f"{col} = :{col}" for col in updates.keys()])
        sql = f"UPDATE {table_name} SET {set_clause} WHERE {where}"
        with self.engine.connect() as conn:
            conn.execute(text(sql), updates)
            conn.commit()
        print(f"更新 {table_name} 表记录（条件：{where}）")

    # ==================== 删（Delete） ====================
    def delete(self, table_name: str, where: str):
        sql = f"DELETE FROM {table_name} WHERE {where}"
        with self.engine.connect() as conn:
            conn.execute(text(sql))
            conn.commit()
        print(f"从 {table_name} 删除记录（条件：{where}）")

    # ==================== 备注 ====================
    def add_table_comment(self, table_name, comment):
        # 反引号包裹表名，避免中文/特殊字符出错
        safe_table = f"`{table_name}`"
        # 转义单引号
        safe_comment = comment.replace("'", "''")
        sql = f"ALTER TABLE {safe_table} COMMENT = '{safe_comment}'"
        with self.engine.begin() as conn:
            conn.execute(text(sql))

    def add_column_comment(self, table_name: str, column_name: str, col_type: str, comment: str):
        with self.engine.connect() as conn:
            sql = f"ALTER TABLE {table_name} MODIFY {column_name} {col_type} COMMENT :comment"
            conn.execute(text(sql), {"comment": comment})
            conn.commit()
        print(f"列 {column_name} 已添加备注：{comment}")

    def get_columns(self, table_name: str) -> list[str]:
        """
        返回指定表的所有列名
        """
        sql = f"SHOW COLUMNS FROM `{table_name}`"
        df = pd.read_sql(sql, self.engine)
        return df['Field'].tolist()

    # ==================== 显示所有表 ====================
    def show_tables(self) -> list[str]:
        """
        获取当前数据库下的所有表名
        """
        sql = "SHOW TABLES"
        df = pd.read_sql(sql, self.engine)
        tables = df.iloc[:, 0].tolist()
        print(f"当前数据库({self.database})共有 {len(tables)} 张表: {tables}")
        return tables

    def find_tables_like(self, keyword: str):
        """
        根据关键字模糊搜索表名，返回包含该关键字的所有表名
        """
        sql = text("""
            SELECT table_name 
            FROM information_schema.tables 
            WHERE table_schema = :db 
              AND table_name LIKE :pattern
        """)
        
        with self.engine.connect() as conn:
            result = conn.execute(sql, {"db": self.database, "pattern": f"%{keyword}%"})
            tables = [row[0] for row in result.fetchall()]
        return tables
    
    # ==================== 通用方法 ====================
    def execute_sql(self, sql: str, params: dict = None):
        with self.engine.connect() as conn:
            conn.execute(text(sql), params or {})
            conn.commit()
        print(f"执行 SQL: {sql}")


    def aggregate_for_key(self, row_name, year, table_list, require_year_for_tables):
        """
        row_name: 行名（可以为 None）
        year: '2020年' 或 None
        table_list: ['母鸡_基础表3  人员支出', ...]
        require_year_for_tables: list of booleans, 与 table_list 对应（True 表示该表在列筛选时需要包含 year）

        返回 (row_name, year, total_amount)
        """
        if not table_list:
            return (row_name, year, 0)

        total_sum = 0.0
        with self.engine.connect() as conn:
            for table, need_year in zip(table_list, require_year_for_tables):
                # --- 获取金额列 ---
                if need_year and year:
                    q = text("""
                        SELECT COLUMN_NAME
                        FROM information_schema.columns
                        WHERE table_name = :table
                        AND (
                            COLUMN_NAME LIKE '%金额%' OR
                            COLUMN_NAME LIKE '%费%' OR
                            COLUMN_NAME LIKE '%小计%' OR
                            COLUMN_NAME LIKE '%合计%'
                        )
                        AND COLUMN_NAME LIKE :year_like
                    """)
                    params = {'table': table, 'year_like': f"%{year}%"}
                else:
                    q = text("""
                        SELECT COLUMN_NAME
                        FROM information_schema.columns
                        WHERE table_name = :table
                        AND (
                            COLUMN_NAME LIKE '%金额%' OR
                            COLUMN_NAME LIKE '%费%' OR
                            COLUMN_NAME LIKE '%小计%' OR
                            COLUMN_NAME LIKE '%合计%'
                        )
                    """)
                    params = {'table': table}

                cols = [r[0] for r in conn.execute(q, params)]
                if not cols:
                    print(f"表 {table} 未找到金额列，跳过。")
                    continue

                # --- 自动获取第二列字段名 ---
                q2 = text("""
                    SELECT COLUMN_NAME
                    FROM information_schema.columns
                    WHERE table_name = :table
                    ORDER BY ORDINAL_POSITION
                    LIMIT 1 OFFSET 1
                """)
                second_col = conn.execute(q2, {'table': table}).scalar()
                if not second_col:
                    print(f"表 {table} 无第二列，跳过。")
                    continue

                # --- 计算金额合计 ---
                sum_expr = " + ".join([f"COALESCE(`{c}`,0)" for c in cols])
                if row_name:
                    sql = text(f"""
                        SELECT SUM({sum_expr}) AS total
                        FROM `{table}`
                        WHERE `{second_col}` LIKE :kw
                    """)

                    # 模糊关键词提取：去掉“部门”等字尾再模糊匹配
                    kw = row_name.replace('支出', '').replace('人员', '') if row_name else ''
                    res = conn.execute(sql, {'kw': f"%{kw}%"}).fetchone()
                else:
                    sql = text(f"""
                        SELECT SUM({sum_expr}) AS total
                        FROM `{table}`
                    """)
                    res = conn.execute(sql).fetchone()

                subtotal = float(res[0]) if res and res[0] is not None else 0.0
                if row_name:
                    print(f"表 {table} | 第二列={second_col} | 关键词='{kw}' | 汇总金额={subtotal}")
                else:
                    print(f"表 {table} | 第二列={second_col} | 汇总金额={subtotal}")
                total_sum += subtotal

        print(f"汇总完成：row_name={row_name}, year={year}, total={total_sum}")
        return (row_name, year, total_sum)

class SheetMapper:
    def __init__(self, df: pd.DataFrame):
        """
        初始化：输入原始 DataFrame（带有 '支出类型' 多列）
        """
        # 取第一行作为表头
        header = df.iloc[0]
        self.df = df[1:].reset_index(drop=True)
        self.df.columns = header

        # 找出“支出类型”相关列
        self.row_name_cols = [col for col in self.df.columns if "支出类型" in str(col)]

        # 添加唯一行名列
        self.df["行名"] = self.df[self.row_name_cols].apply(self._make_row_name, axis=1)

        # 初始化字典
        self.data_dict = self._to_dict()

    def _make_row_name(self, row):
        """
        生成唯一行名：去掉重复和空值
        """
        values = [str(v).strip() for v in row if pd.notna(v) and str(v).strip() != ""]
        values = list(dict.fromkeys(values))  # 去重并保留顺序
        return "_".join(values) if values else "未命名行"

    def _to_dict(self):
        """
        将 DataFrame 转换为字典 {行名_列名: 值}
        """
        data_dict = {}
        for idx, row in self.df.iterrows():
            row_name = row["行名"]
            for col in self.df.columns:
                if col in self.row_name_cols or col == "行名" or col == "":
                    continue
                key = f"{row_name}_{col}"
                data_dict[key] = row[col]
        return data_dict

    def get_dict(self):
        """
        返回 (行名_列名 -> 值) 的字典
        """
        return self.data_dict

    def update_from_dict(self, updates: dict):
        """
        根据 updates = {行名_列名: 值} 更新 DataFrame
        """
        for key, value in updates.items():
            *row_parts, col = key.split("_")
            row_name = "_".join(row_parts)

            # 找行
            row_idx = self.df.index[self.df["行名"] == row_name].tolist()
            if not row_idx:
                print(f"未找到行: {row_name}")
                continue

            # 更新值
            self.df.at[row_idx[0], col] = value
            self.data_dict[key] = value  # 同步更新字典
            
    def get_df(self, drop_row_name=False):
        """
        返回当前 DataFrame
        :param drop_row_name: 是否删除内部生成的 '行名' 列
        """
        df_copy = self.df.copy()
        if drop_row_name and "行名" in df_copy.columns:
            df_copy = df_copy.drop(columns=["行名"])
        return df_copy
    
def excel_to_csv_cleaned(excel_file, mysqlwriter, chatbot, history):
    output_dir = "output_csv"
    os.makedirs("output_csv", exist_ok=True)
    # 打开 Excel 文件
    wb = load_workbook(excel_file, data_only=True)
    target_mapper, meta, header_row = None, None, None
    for sheet_name in wb.sheetnames:
        print(f"正在处理 sheet: {sheet_name} ...")
        ws = wb[sheet_name]

        data = ws.values
        data1 = []
        for row in ws.values:
            clean_row = ["" if cell is None else cell for cell in row]
            data1.append(clean_row)
        df_org = pd.DataFrame(data1)
        df_org = df_org.dropna(how="all")
        for merged in ws.merged_cells.ranges:
            min_row, max_row = merged.min_row - 1, merged.max_row - 1
            min_col, max_col = merged.min_col - 1, merged.max_col - 1
            value = df_org.iat[min_row, min_col]
            for r in range(min_row, max_row + 1):
                for c in range(min_col, max_col + 1):
                    if df_org.iat[r, c] is None or str(df_org.iat[r, c]).strip() == "":
                        df_org.iat[r, c] = value

        df, header_row = detect_headers(df=df_org, keyword="序号", max_header_rows=5)

        if df is None:
            print("跳过当前 sheet（未检测到表头）")
            if "汇总表" in sheet_name:
                target_mapper, meta, header_row = no_header_sheet(df_org, sheet_name, mysqlwriter)
            continue
        meta_info = extract_meta_info(df_org.iloc[:header_row])
        # 删除完全空行
        df = df.dropna(axis=1, how="all")
        first_col = df.columns[0]
        df = df[~df[first_col].astype(str).str.contains("合计|小计", na=False)]
        df.columns = [
            str(c).strip() if c is not None and str(c).strip() != "" else f"未命名列_{i}"
            for i, c in enumerate(df.columns)
        ]
        df = df[[col for col in df.columns if "未命名列" not in col]]
        print(df.columns)
        new_columns = []
        name_count = {}
        for col in df.columns:
            if col not in name_count:
                name_count[col] = 0
                new_columns.append(col)
            else:
                name_count[col] += 1
                new_columns.append(f"{col}（{name_count[col]}）")  # 使用中文括号
        df.columns = new_columns
        text_cols, num_cols = split_text_num_cols(df)
        for col in num_cols:
            df[col] = pd.to_numeric(df[col], errors='coerce').round(2)

        # 序号递增
        if "序号" in str(first_col):
            df[first_col] = range(1, len(df) + 1)
        else:
            df.insert(0, "序号", range(1, len(df) + 1))

        df = replace_ref_with_nan(df)
        df = df.dropna(how="all", subset=[col for col in df.columns if col != "序号"])
        # output_file = os.path.join(output_dir, f"{sheet_name}.csv")
        # df.to_csv(output_file, index=False, encoding="utf-8-sig")
        # print(f"已保存: {output_file}")
        company = meta_info.get("company", "未知公司")  # 如果没有，就用默认值
        sheet_name = f"{company}_{sheet_name}"
        mysqlwriter.insert_df(df, sheet_name)
        mysqlwriter.add_table_comment(sheet_name, meta_info.get("comment", ""))
    return target_mapper, meta, header_row

def split_text_num_cols(df: pd.DataFrame):
    text_cols, num_cols = [], []
    for col in df.columns:
        series = df[col].dropna().astype(str).str.strip()
        # 过滤掉空值、斜杠或特殊占位符
        series = series[~series.isin(['', '/', '-'])]
        if len(series) == 0:
            text_cols.append(col)
            continue
        # 判断剩余值是否全为数值
        is_num = series.str.match(r'^[+-]?(\d+(\.\d+)?([eE][+-]?\d+)?)$').all()
        if is_num:
            num_cols.append(col)
        else:
            text_cols.append(col)
    print(f"文本列: {text_cols}")
    print(f"数值列: {num_cols}")
    return text_cols, num_cols

def replace_ref_with_nan(df: pd.DataFrame) -> pd.DataFrame:
    """
    检查 DataFrame 中是否有 '#REF!'，如果有则替换为 np.nan。
    """
    def clean_cell(x):
        if x == '#REF!':
            return np.nan
        return x

    return df.map(clean_cell)

def detect_headers(df, keyword="序号", max_header_rows=5):
    """
    自动检测并合并多行表头
    参数:
        df: DataFrame
        keyword: 表头识别关键词 (默认找 "序号")
        max_header_rows: 最大表头行数 (默认5)
    返回:
        (处理后的df, 表头行数)
    """
    header_row = None
    # 先找到含有关键字的那一行
    for i in range(min(10, len(df))):
        if df.iloc[i].astype(str).str.contains(keyword).any():
            header_row = i
            break

    if header_row is None:
        return None, None

    # 初始表头
    headers = [df.iloc[header_row].astype(str).tolist()]
    next_row = header_row + 1

    while next_row < len(df) and len(headers) < max_header_rows:
        row = df.iloc[next_row].astype(str)
        first_cell = row.iloc[0].strip()

        if first_cell.isdigit():
            break
        elif first_cell in ["合计", "小计", "总计"]:
            next_row += 1
            continue
        else:
            headers.append(row.tolist())
            next_row += 1

    # 合并多行表头
    final_header = []
    for col_values in zip(*headers):
        col_values = [v for v in col_values if v not in ["nan", "None", ""]]
        cleaned = []
        for v in col_values:
            if not cleaned or cleaned[-1] != v:
                cleaned.append(v)
        final_header.append("_".join(cleaned) if len(cleaned) > 1 else (cleaned[0] if cleaned else ""))

    # 设置列名
    df.columns = final_header

    # 如果是 MultiIndex，也转换成字符串
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = ["_".join([str(c) for c in col if c not in ["nan","None",""]]) 
                      for col in df.columns]
    else:
        df.columns = [str(c) for c in df.columns]

    # 去掉表头行
    df = df.iloc[header_row + len(headers):].reset_index(drop=True)
    df.columns = [str(c).strip() for c in df.columns]

    # 删除“列名为空且该列全为空”的列
    cols_to_keep = []
    for col in df.columns:
        # 如果 df[col] 返回 DataFrame（重复列名情况），取第一列
        series = df[col]
        if isinstance(series, pd.DataFrame):
            series = series.iloc[:, 0]

        # 保留规则：列名非空，或者列里有数据
        if col not in ['', 'None'] or not series.isna().all():
            cols_to_keep.append(col)

    df = df[cols_to_keep].copy()
    return df, header_row

def extract_meta_info(df):
    """
    从表头前的行提取公司名称和comment
    """
    meta = {"company": None, "comment": None}
    rows = df.astype(str).fillna("").iloc[:, 1].tolist()  

    for row in rows:
        row = row.strip()
        if not row:
            continue

        # 提取公司名称
        if "公司名称" in row:
            meta["company"] = row.split("：", 1)[-1].strip()

        # 提取 comment
        if "填表说明" in row:
            meta["comment"] = row
    
    return meta

def no_header_sheet(df_org, sheet_name, mysqlwriter):
    output_flag = "支出汇总表" in sheet_name

    header_row = None
    for i in range(len(df_org)):
        first_col_value = str(df_org.iloc[i, 0]).strip()
        if "支出类型" in first_col_value:
            header_row = i
            break
    if header_row is None:
        print(f"未找到表头：{sheet_name}")
        return None

    meta = extract_meta_info(df_org.iloc[:header_row])
    
    df = df_org.iloc[header_row:].reset_index(drop=True)
    df = df.dropna(axis=1, how="all")
    mapper = SheetMapper(df)
    company = meta.get("company", "未知公司")  # 如果没有，就用默认值
    sheet_name = f"{company}_{sheet_name}"
    mysqlwriter.insert_df(df, sheet_name)
    mysqlwriter.add_table_comment(sheet_name, meta.get("comment", ""))
    if output_flag:
        return mapper, meta, header_row
    else:
        return None, None, None # 返回处理后的 DataFrame
    
@CatchException
def 数据上传(main_input, llm_kwargs, plugin_kwargs, chatbot, history, system_prompt, user_request):
    chatbot.append(["启动数据上传任务", "读取文件并自动处理sheet中..."])
    yield from update_ui(chatbot=chatbot, history=history)
    if os.path.isdir(main_input):
        # 取目录下第一个 Excel 文件
        files = glob.glob(os.path.join(main_input, "*.xlsx"))
        if not files:
            chatbot.append(["错误", "目录中没有找到任何 .xlsx 文件"])
            yield from update_ui(chatbot=chatbot, history=history)
            return
        excel_file = files[0]  # 只取第一个文件
    else:
        if not os.path.isfile(main_input):
            chatbot.append(["错误", f"文件不存在：{main_input}"])
            yield from update_ui(chatbot=chatbot, history=history)
            return
        excel_file = main_input  # main_input 本身就是文件
    
    chatbot.append(["读取文件", os.path.basename(excel_file)])
    yield from update_ui(chatbot=chatbot, history=history)

    total_dict = {}
    mysqlwriter = MySQLWriter("lifeng")
    target_mapper, meta, header_row = excel_to_csv_cleaned(excel_file, mysqlwriter, chatbot, history)
    chatbot.append(["数据上传", "已完成"])
    yield from update_ui(chatbot=chatbot, history=history)
